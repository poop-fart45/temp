from typing import Dict, List
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pymupdf
from ..price_analyzer.analyzer import PriceAnalyzer


class ReportGenerator:
    def __init__(self, quote_data: Dict, output_dir: str = 'media/reports'):
        self.quote_data = quote_data
        self.output_dir = output_dir
        self.price_analyzer = PriceAnalyzer()
        
        # Ensure output directory exists
        os.makedirs(output_dir, exist_ok=True)
        os.makedirs(os.path.join(output_dir, 'pdf'), exist_ok=True)
        os.makedirs(os.path.join(output_dir, 'docx'), exist_ok=True)

    def _format_currency(self, value: float) -> str:
        """Format float value as currency string."""
        return f"${value:,.2f}" if value is not None else "N/A"

    def _format_percentage(self, value: float) -> str:
        """Format float value as percentage string."""
        return f"{value:+.1f}%" if value is not None else "N/A"

    def generate_docx(self) -> str:
        """Generate a DOCX report with quote analysis."""
        doc = Document()
        
        # Add title
        title = doc.add_heading('Quote Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add quote information
        doc.add_heading('Quote Information', level=1)
        doc.add_paragraph(f"Supplier: {self.quote_data['supplier_name']}")
        doc.add_paragraph(f"Quote Number: {self.quote_data['quote_number']}")
        doc.add_paragraph(f"Date: {self.quote_data['quote_date']}")
        
        # Add items table
        doc.add_heading('Items', level=1)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Item Number'
        header_cells[1].text = 'Description'
        header_cells[2].text = 'Quantity'
        header_cells[3].text = 'Unit Price'
        header_cells[4].text = 'Total'
        
        # Add items
        for item in self.quote_data['items']:
            row_cells = table.add_row().cells
            row_cells[0].text = item['item_number']
            row_cells[1].text = item['description']
            row_cells[2].text = str(item['quantity'])
            row_cells[3].text = self._format_currency(item['unit_price'])
            row_cells[4].text = self._format_currency(item['quantity'] * item['unit_price'])
            
            # Add price analysis for each item
            analysis = self.price_analyzer.analyze_item_prices(item['item_number'])
            
            if analysis['has_data']:
                stats = analysis['statistics']
                doc.add_paragraph()
                doc.add_heading(f"Price Analysis - {item['item_number']}", level=2)
                
                p = doc.add_paragraph()
                p.add_run("Historical Price Statistics:\n").bold = True
                p.add_run(f"Minimum Price: {self._format_currency(stats['min_price'])}\n")
                p.add_run(f"Maximum Price: {self._format_currency(stats['max_price'])}\n")
                p.add_run(f"Average Price: {self._format_currency(stats['avg_price'])}\n")
                p.add_run(f"Price Volatility: {self._format_currency(stats['price_volatility'])}\n")
                
                if stats['recent_trend'] is not None:
                    p.add_run(f"90-Day Trend: {self._format_percentage(stats['recent_trend'])}\n")
                
                # Add price trend graph if available
                if analysis['plot_path']:
                    doc.add_picture(analysis['plot_path'], width=Inches(6))
        
        # Save the document
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = os.path.join(self.output_dir, 'docx', f'quote_analysis_{timestamp}.docx')
        doc.save(output_path)
        return output_path

    def generate_pdf(self) -> str:
        """Generate a PDF report with quote analysis."""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        pdf_path = os.path.join(self.output_dir, 'pdf', f'quote_analysis_{timestamp}.pdf')
        
        try:
            # Create a new PDF document
            pdf_doc = pymupdf.open()
            
            # Add a new page
            page = pdf_doc.new_page(width=612, height=792)  # Standard US Letter size
            
            # Set up font sizes
            title_size = 24
            heading_size = 16
            normal_size = 11
            
            # Add title
            y_position = 50
            page.insert_text((306, y_position), "Quote Analysis Report",
                           fontsize=title_size, align=1)  # align=1 centers the text
            
            # Add quote information
            y_position += 50
            page.insert_text((50, y_position), "Quote Information",
                           fontsize=heading_size)
            
            y_position += 30
            info_text = f"""Supplier: {self.quote_data['supplier_name']}
Quote Number: {self.quote_data['quote_number']}
Date: {self.quote_data['quote_date']}"""
            
            for line in info_text.split('\n'):
                page.insert_text((50, y_position), line, fontsize=normal_size)
                y_position += 20
            
            # Add items section
            y_position += 20
            page.insert_text((50, y_position), "Items", fontsize=heading_size)
            y_position += 30
            
            # Process each item
            for item in self.quote_data['items']:
                # Check if we need a new page
                if y_position > 700:
                    page = pdf_doc.new_page(width=612, height=792)
                    y_position = 50
                
                item_text = f"""Item Number: {item['item_number']}
Description: {item['description']}
Quantity: {item['quantity']}
Unit Price: {self._format_currency(item['unit_price'])}
Total: {self._format_currency(item['quantity'] * item['unit_price'])}"""
                
                for line in item_text.split('\n'):
                    page.insert_text((50, y_position), line, fontsize=normal_size)
                    y_position += 20
                
                # Add price analysis if available
                analysis = self.price_analyzer.analyze_item_prices(item['item_number'])
                if analysis['has_data']:
                    stats = analysis['statistics']
                    
                    # Check if we need a new page for the analysis
                    if y_position > 600:  # Leave room for graph
                        page = pdf_doc.new_page(width=612, height=792)
                        y_position = 50
                    
                    y_position += 20
                    page.insert_text((50, y_position),
                                   f"Price Analysis - {item['item_number']}",
                                   fontsize=heading_size)
                    y_position += 30
                    
                    stats_text = f"""Minimum Price: {self._format_currency(stats['min_price'])}
Maximum Price: {self._format_currency(stats['max_price'])}
Average Price: {self._format_currency(stats['avg_price'])}
Price Volatility: {self._format_currency(stats['price_volatility'])}"""
                    
                    if stats['recent_trend'] is not None:
                        stats_text += f"\n90-Day Trend: {self._format_percentage(stats['recent_trend'])}"
                    
                    for line in stats_text.split('\n'):
                        page.insert_text((50, y_position), line, fontsize=normal_size)
                        y_position += 20
                    
                    # Add price trend graph if available
                    if analysis['plot_path']:
                        if y_position > 600:  # Need a new page for the graph
                            page = pdf_doc.new_page(width=612, height=792)
                            y_position = 50
                        
                        # Insert the graph
                        graph_rect = pymupdf.Rect(50, y_position, 550, y_position + 300)
                        page.insert_image(graph_rect, filename=analysis['plot_path'])
                        y_position += 320
                
                y_position += 30  # Space between items
            
            # Save the PDF
            pdf_doc.save(pdf_path)
            pdf_doc.close()
            return pdf_path
            
        except Exception as e:
            print(f"Error creating PDF: {e}")
            return None 